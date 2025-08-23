import os
from docx import Document
from jinja2 import Template
from llm import get_llm
from llm.prompts import TAILOR_PROMPT, COVER_LETTER_PROMPT
from parsers.resume_parser import ResumeProfile
from storage.models import JobPost
from config import cfg
from utils.logger import get_logger

logger = get_logger("tailor")

def _save_docx_paragraphs(paragraphs: list[str], out_path: str):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(out_path)

def _render_resume_template(summary: str, bullets: list[str], skills: str, out_path: str, template_path: str | None = None):
    """Render into a DOCX template preserving styles using placeholders:
    {{SUMMARY}}, {{BULLETS}}, {{SKILLS}}
    - {{BULLETS}} will be replaced by multiple bullet paragraphs using the template's default bullet style.
    """
    try:
        tpl = template_path or cfg.resume_template_path
        if not tpl or not os.path.exists(tpl):
            raise FileNotFoundError("Template not found")
        doc = Document(tpl)
        # Helper: replace text in a paragraph's runs
        def replace_in_paragraph(par, placeholder, replacement):
            if placeholder in par.text:
                # Clear runs and set new text to preserve paragraph style
                for _ in range(len(par.runs)):
                    par.runs[0].clear()
                par.text = par.text.replace(placeholder, replacement)

        # First pass: replace SUMMARY and SKILLS placeholders directly (curly + bracket anchors)
        bullet_anchor_idx = None
        used_placeholder = False
        used_bracket = False
        for i, p in enumerate(doc.paragraphs):
            if "{{SUMMARY}}" in p.text:
                replace_in_paragraph(p, "{{SUMMARY}}", summary)
                used_placeholder = True
            if "{{SKILLS}}" in p.text:
                replace_in_paragraph(p, "{{SKILLS}}", skills)
                used_placeholder = True
            if "{{BULLETS}}" in p.text and bullet_anchor_idx is None:
                bullet_anchor_idx = i
            # Bracket anchors
            if "[TAILOR_SUMMARY]" in p.text:
                replace_in_paragraph(p, "[TAILOR_SUMMARY]", summary)
                used_bracket = True
            if "[TAILOR_SKILLS]" in p.text:
                replace_in_paragraph(p, "[TAILOR_SKILLS]", skills)
                used_bracket = True
            if "[TAILOR_BULLETS]" in p.text and bullet_anchor_idx is None:
                bullet_anchor_idx = i

        # Replace BULLETS anchor with individual bullet paragraphs
        if bullet_anchor_idx is not None:
            anchor_par = doc.paragraphs[bullet_anchor_idx]
            # Replace anchor text with first bullet, then insert others after
            anchor_par.text = ""
            for idx, b in enumerate(bullets):
                par = anchor_par if idx == 0 else doc.add_paragraph()
                par.text = b
                try:
                    par.style = 'List Bullet'
                except Exception:
                    pass
            logger.info("Tailor path: anchor-bullets (%s)", "curly" if used_placeholder else ("bracket" if used_bracket else "unknown"))
        else:
            # No explicit placeholders found; try smart in-place replacement using headers
            lower = lambda s: (s or "").strip().lower()
            n = len(doc.paragraphs)
            def is_header(par):
                name = getattr(par.style, 'name', '') or ''
                txt = par.text.strip()
                return ('heading' in name.lower()) or (len(txt) <= 60 and txt.endswith(':'))

            # Replace Summary section
            for i, p in enumerate(doc.paragraphs):
                txt = lower(p.text)
                if any(k in txt for k in ["summary", "professional summary", "profile", "about"]):
                    # find the next non-empty non-header paragraph to replace
                    j = i + 1
                    while j < n and (not doc.paragraphs[j].text.strip() or is_header(doc.paragraphs[j])):
                        j += 1
                    if j < n:
                        doc.paragraphs[j].text = summary
                    logger.info("Tailor path: smart SUMMARY")
                    break

            # Replace Skills section
            for i, p in enumerate(doc.paragraphs):
                txt = lower(p.text)
                if "skill" in txt:
                    # put skills in the first body paragraph after this header
                    j = i + 1
                    while j < n and is_header(doc.paragraphs[j]):
                        j += 1
                    if j < n:
                        doc.paragraphs[j].text = skills
                    break

            # Replace Highlights/Key Achievements bullets
            for i, p in enumerate(doc.paragraphs):
                txt = lower(p.text)
                if any(k in txt for k in ["highlights", "key achievements", "key highlights", "accomplishments"]):
                    # Overwrite following bullet list items with our bullets
                    j = i + 1
                    bi = 0
                    while j < n and bi < len(bullets):
                        par = doc.paragraphs[j]
                        if is_header(par):
                            break
                        par.text = bullets[bi]
                        try:
                            par.style = 'List Bullet'
                        except Exception:
                            pass
                        bi += 1
                        j += 1
                    # If more bullets remain, append them
                    while bi < len(bullets):
                        par = doc.add_paragraph(bullets[bi])
                        try:
                            par.style = 'List Bullet'
                        except Exception:
                            pass
                        bi += 1
                    logger.info("Tailor path: smart BULLETS (%d)", len(bullets))
                    break
        doc.save(out_path)
    except Exception:
        # Fallback to simple writer if anything fails
        _save_docx_paragraphs(["SUMMARY:", summary, "", "HIGHLIGHTS:"] + [f"- {b}" for b in bullets] + ["", "KEYWORDS:", skills], out_path)

def _render_cover_letter_jinja(context: dict, out_path: str):
    with open(os.path.join(os.path.dirname(__file__), "cover_letter_template.jinja"), "r", encoding="utf-8") as f:
        template = Template(f.read())
    rendered = template.render(**context)
    doc = Document()
    for line in rendered.splitlines():
        doc.add_paragraph(line)
    doc.save(out_path)

def tailor_resume_and_cl(profile: ResumeProfile, job: JobPost, job_slug: str) -> tuple[str, str]:
    llm = get_llm(mode=cfg.llm_mode)
    job_text = job.description or job.title
    # Generate tailored content with richer context
    tailored = llm.generate(
        TAILOR_PROMPT.format(
            job_title=job.title or "",
            company=job.company or "",
            location=job.location or "",
            job_url=job.url or "",
            resume_text=profile.raw_text,
            job_text=job_text,
        ),
        max_tokens=900,
    )
    # parse sections with fallbacks
    summary, bullets, skills = "", [], ""
    import re
    s = tailored or ""
    sm = re.search(r"##\s*SUMMARY(.*?)(##|$)", s, flags=re.S|re.I)
    if sm and sm.group(1).strip():
        summary = sm.group(1).strip()
    bm = re.search(r"##\s*BULLETS(.*?)(##|$)", s, flags=re.S|re.I)
    if bm and bm.group(1).strip():
        bullets = [ln.strip("- ").strip() for ln in bm.group(1).strip().splitlines() if ln.strip()]
    km = re.search(r"##\s*SKILLS(.*?)(##|$)", s, flags=re.S|re.I)
    if km and km.group(1).strip():
        skills = km.group(1).strip()
    # Fallbacks if sections missing
    if not bullets:
        bullets = [ln.strip() for ln in s.splitlines() if ln.strip().startswith('-')][:8]
    if not summary:
        # Take first non-empty line as a minimal summary
        for ln in s.splitlines():
            if ln.strip() and not ln.strip().startswith('-') and not ln.strip().lower().startswith('##'):
                summary = ln.strip()
                break

    # Create a concise tailored resume docx (summary + bullets)
    os.makedirs("output/tailored", exist_ok=True)
    resume_out = f"output/tailored/{job_slug}_resume.docx"
    template_path = None
    if cfg.resume_template_path and os.path.exists(cfg.resume_template_path):
        template_path = cfg.resume_template_path
    elif cfg.resume_path.lower().endswith('.docx') and os.path.exists(cfg.resume_path):
        template_path = cfg.resume_path

    if template_path:
        _render_resume_template(summary, bullets, skills, resume_out, template_path=template_path)
    else:
        _save_docx_paragraphs(["SUMMARY:", summary, "", "HIGHLIGHTS:"] + [f"- {b}" for b in bullets] + ["", "KEYWORDS:", skills], resume_out)

    # Cover letter
    base_text = ""
    if os.path.exists(cfg.cover_letter_base_path):
        from docx import Document as D
        base_text = "\n".join(p.text for p in D(cfg.cover_letter_base_path).paragraphs)

    cl_text = llm.generate(
        COVER_LETTER_PROMPT.format(
            job_title=job.title or "",
            company=job.company or "",
            location=job.location or "",
            job_url=job.url or "",
            name=profile.name or "",
            email=profile.email or "",
            phone=profile.phone or "",
            base_text=base_text,
            job_text=job_text,
        ),
        max_tokens=700,
    )

    cl_out = f"output/tailored/{job_slug}_cover_letter.docx"
    _render_cover_letter_jinja(
        {
            "cover_letter_text": cl_text,
            "name": profile.name or "",
            "email": profile.email or "",
            "phone": profile.phone or "",
            "company": job.company or "Hiring Manager",
            "job_title": job.title or "the role",
        },
        cl_out,
    )
    return resume_out, cl_out
