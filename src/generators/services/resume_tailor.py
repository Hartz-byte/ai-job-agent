import os
import logging
from typing import Optional, Dict, List
from docx import Document

from ..models.tailored_data import TailoredResumeData
from ..utils.docx_utils import (
    add_paragraph_with_style,
    add_section,
    add_bullet_points,
    get_or_create_style
)
from ..utils.template_utils import (
    get_template_path,
    create_document,
    save_document
)
from llm import get_local_llm
from llm.prompts import TAILOR_PROMPT

logger = logging.getLogger("tailor")

class ResumeTailor:
    """Handles the core resume tailoring functionality."""
    
    def __init__(self, profile, job):
        """Initialize with profile and job data."""
        self.profile = profile
        self.job = job
        
    def generate_tailored_content(self) -> TailoredResumeData:
        """Generate tailored resume content using LLM."""
        # Format the prompt with job and resume data
        prompt = TAILOR_PROMPT.format(
            job_title=getattr(self.job, 'title', ''),
            company=getattr(self.job, 'company', ''),
            location=getattr(self.job, 'location', ''),
            job_url=getattr(self.job, 'url', ''),
            resume_text=getattr(self.profile, 'raw_text', ''),
            job_text=getattr(self.job, 'description', '')
        )
        
        # Get LLM response
        llm = get_local_llm()
        response = llm.generate(prompt)
        
        # Parse the response into TailoredResumeData
        return self._parse_llm_response(response)
    
    def _parse_llm_response(self, response: str) -> TailoredResumeData:
        """Parse LLM response into structured TailoredResumeData."""
        from ..utils.parsing_utils import (
            _parse_summary, _parse_experience, _parse_projects,
            _parse_skills, _parse_education, _parse_research_publications
        )
        
        tailored = TailoredResumeData()
        
        # Parse each section
        _parse_summary(response, tailored)
        _parse_experience(response, tailored)
        _parse_projects(response, tailored)
        _parse_skills(response, tailored)
        _parse_education(response, tailored)
        _parse_research_publications(response, tailored)
        
        return tailored
        
    def create_tailored_resume(
        self, 
        tailored_data: TailoredResumeData,
        output_path: str,
        template_path: Optional[str] = None
    ) -> str:
        """Create a tailored resume document."""
        try:
            # Get template path if not provided
            if template_path is None:
                template_path = get_template_path()
            
            logger.debug(f"Using template: {template_path}")
            
            try:
                # Create document from template or scratch
                if template_path and os.path.exists(template_path):
                    doc = create_document(template_path)
                    logger.info("Updating template with tailored data...")
                    try:
                        self._update_template_with_tailored_data(doc, tailored_data)
                    except Exception as update_error:
                        logger.error(f"Error updating template: {update_error}")
                        # Fall back to building from scratch if template update fails
                        logger.info("Falling back to building resume from scratch")
                        doc = self._build_resume_from_scratch(tailored_data)
                else:
                    logger.warning("No template found, creating resume from scratch")
                    doc = self._build_resume_from_scratch(tailored_data)
                
                # Save the document
                return save_document(doc, output_path)
                
            except Exception as e:
                logger.error(f"Error creating tailored DOCX: {e}")
                raise
                
        except Exception as e:
            logger.error(f"Fatal error in create_tailored_resume: {e}")
            raise
    
    def _update_template_with_tailored_data(
        self, 
        doc: Document, 
        tailored_data: TailoredResumeData
    ) -> None:
        """Update existing template with tailored data while preserving original formatting."""
        logger.info("Updating template with tailored data...")
        
        try:
            # First, ensure all required styles exist
            self._ensure_styles_exist(doc)
            
            # Process each section
            sections = [
                ("SUMMARY", self._update_summary_section, tailored_data.summary, True),
                ("EXPERIENCE", self._update_experience_section, tailored_data.experience, True),
                ("PROJECTS", self._update_projects_section, tailored_data.projects, False),
                ("SKILLS", self._update_skills_section, tailored_data.technical_skills, True),
                ("EDUCATION", self._update_education_section, tailored_data.education, True),
                ("RESEARCH PUBLICATIONS", self._update_research_publications, 
                 getattr(tailored_data, 'research_publications', None), False)
            ]
            
            for section_name, update_func, section_data, required in sections:
                if not section_data and required:
                    logger.warning(f"Missing data for required section: {section_name}")
                    continue
                    
                try:
                    logger.info(f"Updating {section_name} section...")
                    
                    # Find or create section heading
                    heading = self._find_section_heading(doc, section_name)
                    if not heading:
                        # Add a new section if it doesn't exist
                        if doc.paragraphs and doc.paragraphs[-1].text.strip():
                            doc.add_paragraph()  # Add spacing
                        heading = doc.add_heading(section_name.title(), level=1)
                        doc.add_paragraph()  # Add spacing after heading
                    
                    # Update the section content
                    update_func(doc, section_data)
                    
                except Exception as e:
                    logger.warning(f"Failed to update {section_name} section: {str(e)}")
                    logger.debug("Error details:", exc_info=True)
        
        except Exception as e:
            logger.error(f"Error updating template: {str(e)}")
            logger.debug("Error details:", exc_info=True)
            raise
            
    def _ensure_styles_exist(self, doc: Document) -> None:
        """Ensure all required styles exist in the document."""
        required_styles = [
            'Title', 'Subtitle', 'Heading 1', 'Heading 2', 'Heading 3',
            'Normal', 'List Bullet', 'List Number', 'Strong', 'Emphasis'
        ]
        
        for style_name in required_styles:
            get_or_create_style(doc, style_name)
            
        # Ensure proper list numbering
        self._ensure_numbering(doc)
    
    def _build_resume_from_scratch(
        self, 
        tailored_data: TailoredResumeData
    ) -> Document:
        """Build complete resume from scratch when no template available."""
        doc = Document()
        
        # Add name and contact info
        if hasattr(self.profile, 'name') and self.profile.name:
            add_paragraph_with_style(doc, self.profile.name, 'Heading 1')
        
        # Add summary
        if tailored_data.summary:
            add_section(doc, "SUMMARY")
            add_paragraph_with_style(doc, tailored_data.summary)
        
        # Add experience
        if tailored_data.experience:
            add_section(doc, "EXPERIENCE")
            for exp in tailored_data.experience:
                self._add_experience_entry(doc, exp)
        
        # Add projects
        if tailored_data.projects:
            add_section(doc, "PROJECTS")
            for proj in tailored_data.projects:
                self._add_project_entry(doc, proj)
        
        # Add skills
        if tailored_data.technical_skills:
            add_section(doc, "SKILLS")
            for category, skills in tailored_data.technical_skills.items():
                add_paragraph_with_style(doc, f"{category}: {', '.join(skills)}")
        
        # Add education
        if tailored_data.education:
            add_section(doc, "EDUCATION")
            for edu in tailored_data.education:
                if isinstance(edu, dict):
                    add_paragraph_with_style(doc, f"{edu.get('degree', '')}, {edu.get('institution', '')} ({edu.get('year', '')})")
                else:
                    add_paragraph_with_style(doc, str(edu))
        
        # Add research publications if they exist
        if tailored_data.research_publications:
            add_section(doc, "RESEARCH & PUBLICATIONS")
            for pub in tailored_data.research_publications:
                add_paragraph_with_style(doc, f"• {pub}")
        
        return doc
    
    def _update_summary_section(self, doc: Document, summary: str) -> None:
        """Update the summary section in the template while preserving formatting."""
        if not summary:
            return
            
        # Find the summary section
        summary_para = self._find_section_heading(doc, "SUMMARY")
        if not summary_para:
            # If no summary section exists, add it after the first paragraph
            if len(doc.paragraphs) > 0:
                # Insert after the first paragraph (contact info)
                doc.paragraphs[0].insert_paragraph_before("SUMMARY", style='Heading 1')
                doc.paragraphs[1].insert_paragraph_before(summary, style='Body Text')
            else:
                doc.add_heading("SUMMARY", level=1)
                doc.add_paragraph(summary, style='Body Text')
            return
            
        # Clear existing content until next heading
        next_para = self._clear_until_next_heading(summary_para)
        
        # Add the new summary with preserved formatting
        if next_para:
            next_para.insert_paragraph_before(summary, style='Body Text')
        else:
            summary_para.insert_paragraph_before(summary, style='Body Text')
            
    def _update_experience_section(self, doc: Document, experience: List[Dict]) -> None:
        """Update the experience section in the template while preserving formatting."""
        if not experience:
            return
            
        exp_heading = self._find_section_heading(doc, "EXPERIENCE")
        if not exp_heading:
            # If no experience section exists, add it
            doc.add_heading("EXPERIENCE", level=1)
            for exp in experience:
                self._add_experience_entry(doc, exp)
            return
            
        # Clear existing content until next section
        next_para = self._clear_until_next_heading(exp_heading)
        
        # Add new experience entries
        for i, exp in enumerate(experience):
            # Add a blank line between entries, but not after the last one
            if i > 0:
                if next_para:
                    next_para.insert_paragraph_before('')
                else:
                    doc.add_paragraph('')
            
            # Add the experience entry
            self._add_experience_entry(doc, exp)
            
        # Ensure there's a blank line after the section
        if next_para and next_para.text.strip() != '':
            next_para.insert_paragraph_before('')
    
    def _add_experience_entry(self, doc: Document, exp: Dict) -> None:
        """Add a single experience entry to the document with proper formatting."""
        # Build job header with title, company, and duration
        job_header = []
        if 'title' in exp and exp['title']:
            job_header.append(exp['title'])
        if 'company' in exp and exp['company']:
            job_header.append(exp['company'])
        if 'location' in exp and exp['location']:
            job_header.append(exp['location'])
            
        # Add the job header with proper formatting
        if job_header:
            # Use a strong style for the job title/company line
            p = doc.add_paragraph(style='Normal')
            p.add_run(' • '.join(job_header)).bold = True
            
            # Add duration on the same line, right-aligned if available
            if 'duration' in exp and exp['duration']:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run('\t')  # Tab to push duration to the right
                p.add_run(exp['duration']).italic = True
        
        # Add bullet points with proper indentation
        bullets = exp.get('bullets', [])
        if bullets:
            for bullet in bullets:
                if bullet.strip():  # Skip empty bullet points
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = 360000  # 0.25 inch in twips
                    p.paragraph_format.first_line_indent = -360000  # Hanging indent
                    p.add_run(bullet.strip())
    
    def _update_skills_section(self, doc: Document, skills: Dict[str, List[str]]) -> None:
        """Update skills section in the document."""
        try:
            # Find skills section or create it
            skills_heading = self._find_section_heading(doc, "SKILLS")
            if not skills_heading:
                skills_heading = doc.add_heading("Skills", level=1)
                doc.add_paragraph()
            
            # Clear existing content until next heading
            self._clear_until_heading(doc, skills_heading, "SKILLS")
            
            if not skills:
                logger.warning("No skills data provided to update skills section")
                return
            
            # Add skills in a clean format
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                    
                # Add category header
                add_paragraph_with_style(doc, category.upper(), 'Heading 2')
                
                # Add skills as comma-separated list
                skills_text = ', '.join(skill_list)
                add_paragraph_with_style(doc, skills_text, 'Normal')
                
                # Add spacing between categories
                doc.add_paragraph()
                
        except Exception as e:
            logger.error(f"Error updating skills section: {str(e)}")
            logger.debug("Error details:", exc_info=True)
            # Fallback to simple text
            try:
                for category, skill_list in skills.items():
                    doc.add_paragraph(f"{category.upper()}:")
                    doc.add_paragraph(", ".join(skill_list))
            except Exception as e:
                logger.error(f"Failed to add fallback skills content: {e}")
    
    def _update_projects_section(self, doc: Document, projects: List[Dict]) -> None:
        """Update the projects section in the template with improved formatting."""
        if not projects:
            return
            
        # Try different possible section headings
        section_names = ["PROJECTS", "PERSONAL PROJECTS", "SIDE PROJECTS"]
        proj_heading = None
        
        for name in section_names:
            proj_heading = self._find_section_heading(doc, name)
            if proj_heading:
                break
        
        if not proj_heading:
            # If no projects section exists, add it
            doc.add_heading("PROJECTS", level=1)
            next_para = None
        else:
            # Clear existing content until next section
            next_para = self._clear_until_next_heading(proj_heading)
        
        # Add projects with consistent formatting
        for i, proj in enumerate(projects):
            # Add a blank line between entries, but not before the first one
            if i > 0 and next_para:
                next_para.insert_paragraph_before('')
            
            # Add project title and details
            if isinstance(proj, dict):
                # Add project name and optional link/date
                title_parts = []
                if proj.get('name'):
                    title_parts.append(proj['name'])
                if proj.get('technologies'):
                    if isinstance(proj['technologies'], list):
                        tech_str = ', '.join(proj['technologies'])
                    else:
                        tech_str = str(proj['technologies'])
                    title_parts.append(f"Technologies: {tech_str}")
                
                title = ' | '.join(title_parts)
                add_paragraph_with_style(doc, title, style='Heading 2')
                
                # Add project description points
                if proj.get('description'):
                    if isinstance(proj['description'], str):
                        # Split by newlines if it's a string
                        points = [p.strip() for p in proj['description'].split('\n') if p.strip()]
                    else:
                        points = proj['description']
                    
                    for point in points:
                        if point.strip():
                            add_bullet_points(doc, [point])
            else:
                # Fallback for string project entries
                add_paragraph_with_style(doc, str(proj), style='Body Text')
        
        # Add a blank line after the section
        if next_para and next_para.text.strip() != '':
            next_para.insert_paragraph_before('')
    
    def _update_education_section(self, doc: Document, education: List) -> None:
        """Update the education section in the template with improved formatting."""
        if not education:
            return
            
        edu_heading = self._find_section_heading(doc, "EDUCATION")
        if not edu_heading:
            # If no education section exists, add it
            doc.add_heading("EDUCATION", level=1)
            next_para = None
        else:
            # Clear existing content until next section
            next_para = self._clear_until_next_heading(edu_heading)
        
        # Add education entries with consistent formatting
        for i, edu in enumerate(education):
            if i > 0 and next_para:
                next_para.insert_paragraph_before('')
            
            # Create a table for each education entry to align degree and date
            from docx.shared import Inches
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            # Create a table with 2 columns (degree and date)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Normal'  # No borders for a clean look
            table.autofit = False
            
            # Set column widths
            table.columns[0].width = Inches(5.0)  # Degree and school
            table.columns[1].width = Inches(1.5)  # Date (right-aligned)
            
            # Add degree and school
            degree_cell = table.cell(0, 0)
            p = degree_cell.paragraphs[0]
            
            # Add degree in bold
            if 'degree' in edu and edu['degree']:
                p.add_run(edu['degree']).bold = True
                
            # Add school name
            if 'school' in edu and edu['school']:
                if p.text:  # Add a space if there's already text
                    p.add_run(', ')
                p.add_run(edu['school'])
                
            # Add location if available
            if 'location' in edu and edu['location']:
                if p.text:  # Add a space if there's already text
                    p.add_run(' | ')
                p.add_run(edu['location']).italic = True
            
            # Add date (right-aligned)
            if 'date' in edu and edu['date']:
                date_cell = table.cell(0, 1)
                date_cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT
                date_cell.paragraphs[0].add_run(edu['date']).italic = True
            
            # Add any additional details (GPA, honors, etc.)
            if 'details' in edu and edu['details']:
                details = edu['details']
                if isinstance(details, str):
                    doc.add_paragraph(details, style='Body Text')
                elif isinstance(details, list):
                    for detail in details:
                        if detail.strip():
                            doc.add_paragraph(detail, style='Body Text')
        
        # Add a blank line after the section
        if next_para and next_para.text.strip() != '':
            next_para.insert_paragraph_before('')
    
    def _update_research_publications(self, doc: Document, publications: List[str]) -> None:
        """Update the research publications section in the template with improved formatting."""
        if not publications:
            return
            
        # Try different possible section headings
        section_names = ["RESEARCH & PUBLICATIONS", "PUBLICATIONS", "RESEARCH"]
        pub_heading = None
        
        for name in section_names:
            pub_heading = self._find_section_heading(doc, name)
            if pub_heading:
                break
        
        if not pub_heading:
            # If no publications section exists, add it
            doc.add_heading("RESEARCH & PUBLICATIONS", level=1)
            next_para = None
        else:
            # Clear existing content until next section
            next_para = self._clear_until_next_heading(pub_heading)
        
        # Add publications with proper formatting
        for i, pub in enumerate(publications):
            if not pub.strip():
                continue
                
            # Add a blank line between publications, but not before the first one
            if i > 0 and next_para:
                next_para.insert_paragraph_before('')
            
            # Create a paragraph for the publication
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = 360000  # 0.25 inch in twips
            p.paragraph_format.first_line_indent = -360000  # Hanging indent
            
            # Add publication text with proper formatting
            p.add_run(pub.strip())
        
        # Add a blank line after the section
        if next_para and next_para.text.strip() != '':
            next_para.insert_paragraph_before('')
    
    def _find_section_heading(self, doc: Document, section_name: str):
        """Find a section heading in the document."""
        section_name = section_name.upper()
        for para in doc.paragraphs:
            # Check if paragraph text matches the section name (case-insensitive)
            # Also check if it's a heading style
            if (para.text.upper() == section_name or 
                (hasattr(para, 'style') and 
                 para.style.name.startswith('Heading') and 
                 para.text.upper() == section_name)):
                return para
        return None
    
    def _clear_until_heading(self, doc: Document, start_para, target_heading: str):
        """
        Clear content from start_para until the next heading.
        
        Args:
            doc: The document
            start_para: The starting paragraph
            target_heading: The heading text to find (case-insensitive)
            
        Returns:
            The next heading paragraph, or None if not found
        """
        try:
            if not hasattr(start_para, '_element'):
                return None
                
            # Get the parent element containing paragraphs
            parent = start_para._element.getparent()
            if parent is None:
                return None
                
            # Get all paragraphs
            all_paras = list(parent.iterchildren('w:p'))
            if not all_paras:
                return None
                
            # Find the starting index
            start_idx = None
            for i, p in enumerate(all_paras):
                if p == start_para._element:
                    start_idx = i
                    break
                    
            if start_idx is None:
                return None
                
            # Find the next heading
            next_heading = None
            for i in range(start_idx + 1, len(all_paras)):
                p = all_paras[i]
                
                # Check if this is a heading
                is_heading = False
                style_name = p.xpath('.//w:pStyle/@w:val')
                if style_name:
                    style_name = style_name[0].lower()
                    is_heading = style_name.startswith('heading') or style_name in ['title', 'subtitle']
                
                # Check if this is our target heading
                para_text = ''.join(node.text for node in p.xpath('.//w:t') if node.text)
                if is_heading and target_heading.lower() in para_text.lower():
                    next_heading = p
                    break
                    
            # Clear content between start_para and next_heading
            end_idx = all_paras.index(next_heading) if next_heading else len(all_paras)
            
            # Keep the start_para but clear its content
            if start_para._element.text:
                start_para.clear()
                
            # Remove paragraphs between start and end
            for i in range(start_idx + 1, end_idx):
                if i < len(all_paras):
                    try:
                        p = all_paras[i]
                        p.getparent().remove(p)
                    except Exception as e:
                        logger.warning(f"Error removing paragraph: {e}")
                        continue
            
            return next_heading
            
        except Exception as e:
            logger.warning(f"Error in _clear_until_heading: {e}")
            return None
