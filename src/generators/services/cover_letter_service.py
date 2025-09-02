import logging
from datetime import datetime
from typing import Optional, List
from docx import Document

from ..utils.docx_utils import (
    add_paragraph_with_style
)
from ..utils.template_utils import (
    get_template_path,
    create_document,
    save_document
)

logger = logging.getLogger("tailor")

class CoverLetterService:
    """Handles cover letter generation and customization."""
    
    def __init__(self, profile, job):
        """Initialize with profile and job data."""
        self.profile = profile
        self.job = job
    
    def generate_cover_letter(
        self,
        output_path: str,
        template_path: Optional[str] = None
    ) -> str:
        """Generate a tailored cover letter."""
        try:
            # Get template path if not provided
            if template_path is None:
                template_path = get_template_path('cover_letter')
            
            # Create document from template or scratch
            doc = create_document(template_path) if template_path else Document()
            
            # Add content to the cover letter
            self._add_cover_letter_content(doc)
            
            # Save the document
            return save_document(doc, output_path)
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            raise
    
    def _add_cover_letter_content(self, doc: Document) -> None:
        """Add content to the cover letter document."""
        # Add date
        today = datetime.now().strftime("%B %d, %Y")
        add_paragraph_with_style(doc, today, style_name='Normal')
        
        # Add recipient info
        if hasattr(self.job, 'company') and self.job.company:
            add_paragraph_with_style(doc, self.job.company, style_name='Normal')
            add_paragraph_with_style(doc, "Hiring Manager", style_name='Normal')
            add_paragraph_with_style(doc, self.job.company, style_name='Normal')
            if hasattr(self.job, 'location') and self.job.location:
                add_paragraph_with_style(doc, self.job.location, style_name='Normal')
        
        # Add salutation
        salutation = "Dear Hiring Manager,"
        if hasattr(self.job, 'hiring_manager') and self.job.hiring_manager:
            salutation = f"Dear {self.job.hiring_manager},"
        add_paragraph_with_style(doc, salutation, style_name='Normal')
        add_paragraph_with_style(doc, "", style_name='Normal')  # Empty line
        
        # Add introduction
        intro = self._generate_introduction()
        add_paragraph_with_style(doc, intro, style_name='Normal')
        
        # Add body paragraphs
        body_paragraphs = self._generate_body_paragraphs()
        for para in body_paragraphs:
            add_paragraph_with_style(doc, para, style_name='Normal')
        
        # Add closing
        closing = self._generate_closing()
        add_paragraph_with_style(doc, closing, style_name='Normal')
        add_paragraph_with_style(doc, "Sincerely,", style_name='Normal')
        
        # Add signature
        if hasattr(self.profile, 'name') and self.profile.name:
            add_paragraph_with_style(doc, self.profile.name, style_name='Normal')
            
            # Add contact info if available
            contact_info = []
            if hasattr(self.profile, 'email') and self.profile.email:
                contact_info.append(self.profile.email)
            if hasattr(self.profile, 'phone') and self.profile.phone:
                contact_info.append(self.profile.phone)
            if hasattr(self.profile, 'linkedin') and self.profile.linkedin:
                contact_info.append(self.profile.linkedin)
            
            if contact_info:
                add_paragraph_with_style(doc, "  ".join(contact_info), style_name='Normal')
    
    def _generate_introduction(self) -> str:
        """Generate the introduction paragraph of the cover letter."""
        position = getattr(self.job, 'title', 'this position')
        company = getattr(self.job, 'company', 'your company')
        
        intro = f"I am excited to apply for the {position} position at {company}. "
        
        if hasattr(self.profile, 'current_role') and self.profile.current_role:
            intro += f"With my experience as a {self.profile.current_role}, "
            intro += "I am confident in my ability to contribute effectively to your team. "
        else:
            intro += "I am confident that my skills and experience make me a strong candidate. "
        
        return intro
    
    def _generate_body_paragraphs(self) -> List[str]:
        """Generate the body paragraphs of the cover letter."""
        paragraphs = []
        
        # First paragraph - relevant experience
        exp_para = ""
        if hasattr(self.profile, 'experience') and self.profile.experience:
            exp_para = "In my current role, I have "
            exp_para += ", ".join([exp.get('summary', '') for exp in self.profile.experience[:2]])
            exp_para += ". This experience has equipped me with valuable skills that align well with the requirements for this position."
        
        if exp_para:
            paragraphs.append(exp_para)
        
        # Second paragraph - relevant skills
        skills_para = ""
        if hasattr(self.profile, 'skills') and self.profile.skills:
            skills = self.profile.skills[:5]  # Take top 5 skills
            skills_para = f"My technical expertise includes {', '.join(skills[:-1])}, and {skills[-1]}. "
            
            if hasattr(self.job, 'requirements') and self.job.requirements:
                skills_para += "I am particularly drawn to this opportunity because my background in these areas directly aligns with the key requirements you're seeking. "
            
            skills_para += "I am eager to bring my skills and experience to your team and contribute to your company's success."
        
        if skills_para:
            paragraphs.append(skills_para)
        
        # Third paragraph - why you're interested
        interest_para = "I am particularly interested in this opportunity because "
        if hasattr(self.job, 'company') and self.job.company:
            interest_para += f"I admire {self.job.company}'s "
            if hasattr(self.job, 'company_description') and self.job.company_description:
                interest_para += f"{self.job.company_description.lower()} "
            else:
                interest_para += "work in the industry "
            
            interest_para += "and I am excited about the prospect of contributing to your team. "
        
        interest_para += "I am confident that my background and skills would make me a valuable addition to your organization."
        paragraphs.append(interest_para)
        
        return paragraphs
    
    def _generate_closing(self) -> str:
        """Generate the closing paragraph of the cover letter."""
        closing = "Thank you for considering my application. "
        closing += "I would welcome the opportunity to discuss how my skills and experience align with your needs. "
        closing += "I am available at your earliest convenience for an interview and can be reached at "
        
        if hasattr(self.profile, 'phone') and self.profile.phone:
            closing += f"{self.profile.phone} or "
            
        closing += f"{self.profile.email}. "
        closing += "I look forward to the possibility of contributing to your team."
        
        return closing
