import os
import logging
from typing import Dict, Optional, Any
from docx import Document

from ..utils.docx_utils import (
    add_paragraph_with_style,
    add_section,
    add_bullet_points
)
from ..utils.template_utils import (
    create_document,
    save_document
)

logger = logging.getLogger("tailor")

class FallbackService:
    """Handles fallback mechanisms when primary tailoring fails."""
    
    def __init__(self, profile, job):
        """Initialize with profile and job data."""
        self.profile = profile
        self.job = job
    
    def create_basic_resume(
        self,
        output_path: str,
        template_path: Optional[str] = None
    ) -> str:
        """Create a basic resume when comprehensive tailoring fails."""
        try:
            # Create a new document
            doc = create_document(template_path) if template_path else Document()
            
            # Add basic resume content
            self._add_basic_resume_content(doc)
            
            # Save the document
            return save_document(doc, output_path)
            
        except Exception as e:
            logger.error(f"Error creating basic resume: {e}")
            # Last resort: create a simple text file
            return self._create_simple_text_fallback(output_path)
    
    def _add_basic_resume_content(self, doc: Document) -> None:
        """Add basic resume content to the document."""
        # Add name and contact info
        if hasattr(self.profile, 'name') and self.profile.name:
            add_paragraph_with_style(doc, self.profile.name, 'Heading 1')
            
            # Add contact information
            contact_info = []
            if hasattr(self.profile, 'email') and self.profile.email:
                contact_info.append(self.profile.email)
            if hasattr(self.profile, 'phone') and self.profile.phone:
                contact_info.append(self.profile.phone)
            if hasattr(self.profile, 'linkedin') and self.profile.linkedin:
                contact_info.append(self.profile.linkedin)
            if hasattr(self.profile, 'location') and self.profile.location:
                contact_info.append(self.profile.location)
                
            if contact_info:
                add_paragraph_with_style(doc, " | ".join(contact_info), 'Normal')
        
        # Add summary if available
        if hasattr(self.profile, 'summary') and self.profile.summary:
            add_section(doc, "SUMMARY")
            add_paragraph_with_style(doc, self.profile.summary, 'Normal')
        
        # Add experience
        if hasattr(self.profile, 'experience') and self.profile.experience:
            add_section(doc, "EXPERIENCE")
            for exp in self.profile.experience[:3]:  # Limit to 3 most recent
                self._add_basic_experience_entry(doc, exp)
        
        # Add skills
        if hasattr(self.profile, 'skills') and self.profile.skills:
            add_section(doc, "SKILLS")
            # Group skills if they're categorized
            if isinstance(self.profile.skills, dict):
                for category, skills in self.profile.skills.items():
                    if skills:
                        add_paragraph_with_style(
                            doc,
                            f"{category}: {', '.join(skills[:10])}",
                            'Normal'
                        )
            else:
                # Just a flat list of skills
                add_paragraph_with_style(
                    doc,
                    ', '.join(self.profile.skills[:20]),  # Limit to top 20 skills
                    'Normal'
                )
        
        # Add education
        if hasattr(self.profile, 'education') and self.profile.education:
            add_section(doc, "EDUCATION")
            for edu in self.profile.education[:2]:  # Limit to 2 most recent
                self._add_basic_education_entry(doc, edu)
    
    def _add_basic_experience_entry(self, doc: Document, exp: Dict[str, Any]) -> None:
        """Add a basic experience entry to the document."""
        # Build the job header
        header_parts = []
        if 'title' in exp and exp['title']:
            header_parts.append(exp['title'])
        if 'company' in exp and exp['company']:
            header_parts.append(exp['company'])
        if 'dates' in exp and exp['dates']:
            header_parts.append(exp['dates'])
        
        if header_parts:
            add_paragraph_with_style(doc, " | ".join(header_parts), 'Heading 3')
        
        # Add bullet points if available
        if 'highlights' in exp and exp['highlights']:
            add_bullet_points(doc, exp['highlights'][:5])  # Limit to 5 bullet points
    
    def _add_basic_education_entry(self, doc: Document, edu) -> None:
        """Add a basic education entry to the document.
        
        Args:
            doc: The document to add to
            edu: Either an Education object or a dictionary with education data
        """
        edu_parts = []
        
        # Handle both Education objects and dictionaries
        if hasattr(edu, 'degree') and edu.degree:
            edu_parts.append(edu.degree)
        elif isinstance(edu, dict) and 'degree' in edu and edu['degree']:
            edu_parts.append(edu['degree'])
            
        if hasattr(edu, 'institution') and edu.institution:
            edu_parts.append(edu.institution)
        elif isinstance(edu, dict) and 'institution' in edu and edu['institution']:
            edu_parts.append(edu['institution'])
            
        if hasattr(edu, 'year') and edu.year:
            edu_parts.append(f"({edu.year})")
        elif hasattr(edu, 'dates') and edu.dates:
            edu_parts.append(f"({edu.dates})")
        elif isinstance(edu, dict) and 'dates' in edu and edu['dates']:
            edu_parts.append(f"({edu['dates']})")
        elif isinstance(edu, dict) and 'year' in edu and edu['year']:
            edu_parts.append(f"({edu['year']})")
        
        if edu_parts:
            add_paragraph_with_style(doc, " ".join(edu_parts), 'Normal')
    
    def _create_simple_text_fallback(self, output_path: str) -> str:
        """Create a simple text file as a last resort fallback."""
        try:
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True, mode=0o755)
            
            # Create a simple text version
            text_content = []
            
            # Add name and contact info
            if hasattr(self.profile, 'name') and self.profile.name:
                text_content.append(self.profile.name.upper())
                text_content.append("=" * len(self.profile.name))
                text_content.append("")
                
                # Add contact information
                contact_info = []
                if hasattr(self.profile, 'email') and self.profile.email:
                    contact_info.append(f"Email: {self.profile.email}")
                if hasattr(self.profile, 'phone') and self.profile.phone:
                    contact_info.append(f"Phone: {self.profile.phone}")
                if hasattr(self.profile, 'linkedin') and self.profile.linkedin:
                    contact_info.append(f"LinkedIn: {self.profile.linkedin}")
                if hasattr(self.profile, 'location') and self.profile.location:
                    contact_info.append(f"Location: {self.profile.location}")
                
                if contact_info:
                    text_content.append(" | ".join(contact_info))
                    text_content.append("")
            
            # Add summary
            if hasattr(self.profile, 'summary') and self.profile.summary:
                text_content.append("SUMMARY")
                text_content.append("-" * 7)
                text_content.append(self.profile.summary)
                text_content.append("")
            
            # Add experience
            if hasattr(self.profile, 'experience') and self.profile.experience:
                text_content.append("EXPERIENCE")
                text_content.append("-" * 10)
                for exp in self.profile.experience[:3]:  # Limit to 3 most recent
                    # Add job header
                    header_parts = []
                    if 'title' in exp and exp['title']:
                        header_parts.append(exp['title'])
                    if 'company' in exp and exp['company']:
                        header_parts.append(exp['company'])
                    if 'dates' in exp and exp['dates']:
                        header_parts.append(exp['dates'])
                    
                    if header_parts:
                        text_content.append(" | ".join(header_parts))
                    
                    # Add bullet points
                    if 'highlights' in exp and exp['highlights']:
                        for point in exp['highlights'][:3]:  # Limit to 3 bullet points
                            text_content.append(f"- {point}")
                    
                    text_content.append("")  # Empty line between entries
            
            # Add skills
            if hasattr(self.profile, 'skills') and self.profile.skills:
                text_content.append("SKILLS")
                text_content.append("-" * 6)
                
                if isinstance(self.profile.skills, dict):
                    for category, skills in self.profile.skills.items():
                        if skills:
                            text_content.append(f"{category}: {', '.join(skills[:10])}")
                else:
                    # Just a flat list of skills
                    text_content.append(", ".join(self.profile.skills[:20]))  # Limit to top 20 skills
                
                text_content.append("")  # Empty line
            
            # Add education
            if hasattr(self.profile, 'education') and self.profile.education:
                text_content.append("EDUCATION")
                text_content.append("-" * 9)
                for edu in self.profile.education[:2]:  # Limit to 2 most recent
                    edu_parts = []
                    if 'degree' in edu and edu['degree']:
                        edu_parts.append(edu['degree'])
                    if 'institution' in edu and edu['institution']:
                        edu_parts.append(f"at {edu['institution']}")
                    if 'dates' in edu and edu['dates']:
                        edu_parts.append(f"({edu['dates']})")
                    
                    if edu_parts:
                        text_content.append(" ".join(edu_parts))
                
                text_content.append("")  # Empty line
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(text_content))
            
            logger.info(f"Created simple text resume at {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to create simple text fallback: {e}")
            raise
