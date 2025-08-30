import logging
from typing import List
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("tailor")

def get_or_create_style(doc: Document, style_name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Get existing style or create it if it doesn't exist, with fallback to Normal."""
    try:
        # First try to get the style directly
        if style_name in doc.styles:
            return doc.styles[style_name]
            
        # Try common variations of the style name
        variations = [
            style_name,
            style_name.lower(),
            style_name.upper(),
            style_name.title(),
        ]
        
        for variation in variations:
            if variation in doc.styles:
                return doc.styles[variation]
                
        # If style doesn't exist, create it
        return _create_style(doc, style_name, style_type)
        
    except Exception as e:
        logger.warning(f"Error getting style '{style_name}': {e}")
        return doc.styles.get('Normal', doc.styles['Default Paragraph Font'])

def _create_style(doc: Document, style_name: str, style_type):
    """Create a new style with the given name and type."""
    try:
        # Check if style already exists (case-insensitive)
        for existing_style in doc.styles:
            if existing_style.name.lower() == style_name.lower():
                return existing_style
        
        # Create new style
        style = doc.styles.add_style(style_name, style_type)
        
        # Set font properties
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Configure style based on name
        style_config = {
            'heading 1': {'size': 16, 'bold': True, 'space_after': Pt(12)},
            'heading 2': {'size': 14, 'bold': True, 'space_after': Pt(6)},
            'heading 3': {'size': 12, 'bold': True, 'italic': True, 'space_after': Pt(6)},
            'title': {'size': 18, 'bold': True, 'space_after': Pt(18)},
            'subtitle': {'size': 14, 'italic': True, 'space_after': Pt(12)},
            'list bullet': {'size': 11, 'space_before': Pt(0), 'space_after': Pt(0)},
            'list': {'size': 11, 'space_before': Pt(0), 'space_after': Pt(0)},
            'normal': {'size': 11, 'space_after': Pt(6)},
        }
        
        # Apply style configuration
        config = style_config.get(style_name.lower(), style_config['normal'])
        
        if 'size' in config:
            font.size = Pt(config['size'])
        if 'bold' in config:
            font.bold = config['bold']
        if 'italic' in config:
            font.italic = config['italic']
        if 'space_before' in config:
            style.paragraph_format.space_before = config['space_before']
        if 'space_after' in config:
            style.paragraph_format.space_after = config['space_after']
        
        # For list styles, ensure proper indentation
        if 'list' in style_name.lower():
            style.paragraph_format.left_indent = Inches(0.25)
            style.paragraph_format.first_line_indent = Inches(-0.25)
        
        return style
        
    except Exception as e:
        logger.warning(f"Could not create style '{style_name}': {e}")
        # Return Normal style as fallback
        return doc.styles.get('Normal', None)

def add_paragraph_with_style(doc: Document, text: str, style_name: str = None, **kwargs):
    """
    Add paragraph with style, handling missing styles gracefully.
    
    Args:
        doc: The document to add the paragraph to
        text: The text content of the paragraph
        style_name: The name of the style to apply
        **kwargs: Additional formatting options (bold, italic, etc.)
        
    Returns:
        The created paragraph
    """
    try:
        # Handle empty or whitespace-only text
        if not text or not str(text).strip():
            return doc.add_paragraph()
            
        # Get or create the style
        style = None
        if style_name:
            style = get_or_create_style(doc, style_name)
        
        # Add the paragraph with the specified style
        para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        
        # Add runs with formatting
        run = para.add_run(text)
        
        # Apply formatting from kwargs
        if 'bold' in kwargs:
            run.bold = kwargs['bold']
        if 'italic' in kwargs:
            run.italic = kwargs['italic']
        if 'underline' in kwargs:
            run.underline = kwargs['underline']
        if 'font_size' in kwargs:
            run.font.size = Pt(kwargs['font_size'])
        if 'font_name' in kwargs:
            run.font.name = kwargs['font_name']
        if 'color' in kwargs:
            run.font.color.rgb = kwargs['color']
            
        return para
        
    except Exception as e:
        logger.error(f"Error adding paragraph with style '{style_name}': {e}")
        # Fallback to basic paragraph
        try:
            return doc.add_paragraph(text)
        except Exception as e:
            logger.error(f"Failed to add fallback paragraph: {e}")
            return doc.add_paragraph("Error: Could not add paragraph")

def add_section(doc: Document, title: str, level: int = 2):
    """
    Add a new section to the document with proper spacing.
    
    Args:
        doc: The document to add the section to
        title: The section title text
        level: The heading level (1-3)
    """
    try:
        if doc.paragraphs and doc.paragraphs[-1].text.strip() != '':
            doc.add_paragraph('')
        
        heading = doc.add_heading(title, level=min(max(1, level), 3))
        doc.add_paragraph('')
        return heading
    except Exception as e:
        logger.error(f"Error adding section '{title}': {e}")
        doc.add_paragraph(title, style='Heading 2')
        doc.add_paragraph('')

def add_bullet_points(doc: Document, items: List[str], style_name: str = 'List Bullet'):
    """Add bullet points to the document."""
    for item in items:
        para = add_paragraph_with_style(doc, item, style_name)
        _set_list_style(para)

def _set_list_style(paragraph, level=0):
    """
    Set the list style for a paragraph with proper numbering.
    
    Args:
        paragraph: The paragraph to format as a list item
        level: The indentation level (0 for top-level, 1 for nested, etc.)
    """
    try:
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        
        # Remove any existing numbering
        for numPr in pPr.xpath('.//w:numPr'):
            pPr.remove(numPr)
        
        # Create new numbering properties
        numPr = OxmlElement('w:numPr')
        
        # Set indentation level
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(min(level, 8)))  # Max 8 levels deep
        
        # Set numbering ID (1 for bullet, 2 for numbered)
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')  # 1 for bullet, 2 for numbered
        
        # Build the numbering properties
        numPr.append(ilvl)
        numPr.append(numId)
        
        # Add to paragraph properties
        pPr.append(numPr)
        
        # Set proper indentation
        pPr_ind = pPr.get_or_add_ind()
        pPr_ind.left = qn('w:left')
        pPr_ind.left = str(int(level * 720))  # 720 twips = 0.5"
        pPr_ind.hanging = qn('w:hanging')
        pPr_ind.hanging = '360'  # 360 twips = 0.25"
        
    except Exception as e:
        logger.warning(f"Error setting list style: {e}")
        # Fallback: Just add a bullet character
        if paragraph.text and not paragraph.text.startswith('•'):
            paragraph.text = '• ' + paragraph.text
