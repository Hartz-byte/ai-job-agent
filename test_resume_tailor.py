import sys
import os
import logging
from datetime import datetime
from pathlib import Path

# Add src directory to Python path
sys.path.insert(0, os.path.abspath('.'))

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("test")

# Create output directory if it doesn't exist
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

def create_test_job():
    """Create a test job with sample requirements."""
    class TestJob:
        def __init__(self):
            self.title = "Senior Software Engineer"
            self.company = "Tech Innovations Inc."
            self.location = "San Francisco, CA (Hybrid)"
            self.url = "https://example.com/job/senior-software-engineer-123"
            self.description = """
            We are looking for an experienced Senior Software Engineer to join our team.
            
            Key Responsibilities:
            - Design and develop scalable web applications using Python and modern frameworks
            - Lead technical architecture decisions and mentor junior developers
            - Implement best practices for code quality, testing, and deployment
            - Collaborate with cross-functional teams to deliver high-quality software
            
            Requirements:
            - 5+ years of professional software development experience
            - Strong expertise in Python and web frameworks (Django, FastAPI, or Flask)
            - Experience with cloud platforms (AWS, GCP, or Azure)
            - Knowledge of containerization and orchestration (Docker, Kubernetes)
            - Experience with databases (PostgreSQL, MongoDB)
            - Strong problem-solving and communication skills
            
            Nice to Have:
            - Experience with machine learning and data pipelines
            - Knowledge of frontend technologies (React, TypeScript)
            - Experience with CI/CD pipelines
            - Open source contributions
            """
    return TestJob()

def generate_test_resume():
    """Generate a test resume document with sample content."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    
    doc = Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add header with name and contact info
    header = doc.add_paragraph()
    header_run = header.add_run("JOHN DOE")
    header_run.bold = True
    header_run.font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("San Francisco, CA • johndoe@email.com • (123) 456-7890 • linkedin.com/in/johndoe")
    
    # Add summary section
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph("""
    Results-driven Senior Software Engineer with 7+ years of experience in designing and developing 
    scalable web applications. Proven track record of leading development teams and delivering 
    high-impact solutions. Specialized in Python, Django, and cloud technologies.
    """.strip())
    
    # Add experience section
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    
    # Experience 1
    exp1 = doc.add_paragraph()
    exp1.add_run("Senior Software Engineer").bold = True
    exp1.add_run(" | Tech Solutions Inc. | San Francisco, CA | ")
    exp1.add_run("Jan 2020 - Present").italic = True
    doc.add_paragraph("""
    • Led a team of 5 developers in building a scalable microservices architecture serving 1M+ users
    • Architected and implemented RESTful APIs using Django and FastAPI, improving performance by 40%
    • Mentored junior developers and established coding best practices across the team
    • Integrated CI/CD pipelines reducing deployment time by 60%
    """.strip())
    
    # Experience 2
    exp2 = doc.add_paragraph()
    exp2.add_run("Software Engineer").bold = True
    exp2.add_run(" | Digital Innovations | San Jose, CA | ")
    exp2.add_run("Jun 2017 - Dec 2019").italic = True
    doc.add_paragraph("""
    • Developed and maintained RESTful APIs using Django and Flask, handling 100K+ daily requests
    • Designed and optimized database schemas, reducing query times by 65%
    • Implemented automated testing, increasing test coverage from 60% to 95%
    • Collaborated with cross-functional teams to deliver features on time
    """.strip())
    
    # Add skills section
    doc.add_heading('TECHNICAL SKILLS', level=1)
    skills = doc.add_paragraph()
    skills.add_run("Programming Languages: ").bold = True
    skills.add_run("Python, JavaScript, SQL, Java, TypeScript")
    
    # Add education section
    doc.add_heading('EDUCATION', level=1)
    edu = doc.add_paragraph()
    edu.add_run("MS in Computer Science").bold = True
    edu.add_run(" | Stanford University | ")
    edu.add_run("2015 - 2017").italic = True
    
    # Save the test resume
    test_resume_path = OUTPUT_DIR / "test_resume.docx"
    doc.save(test_resume_path)
    return test_resume_path

def test_resume_tailoring():
    """Test the resume tailoring functionality."""
    try:
        logger.info("Starting resume tailoring test...")
        
        # Import required modules
        from src.parsers.resume_parser import parse_resume
        from src.generators.services.resume_tailor import ResumeTailor
        
        # Generate a test resume if it doesn't exist
        test_resume_path = OUTPUT_DIR / "test_resume.docx"
        if not test_resume_path.exists():
            logger.info("Generating test resume...")
            test_resume_path = generate_test_resume()
        
        # Parse the test resume
        logger.info("Parsing test resume...")
        profile = parse_resume(str(test_resume_path))
        
        # Create a test job
        test_job = create_test_job()
        
        # Initialize the resume tailor
        logger.info("Initializing ResumeTailor...")
        tailor = ResumeTailor(profile, test_job)
        
        # Generate tailored content
        logger.info("Generating tailored content...")
        tailored_data = tailor.generate_tailored_content()
        
        # Print tailored content
        logger.info("\n=== TAILORED SUMMARY ===")
        print(tailored_data.summary)
        
        logger.info("\n=== TAILORED SKILLS ===")
        for category, skills in tailored_data.technical_skills.items():
            print(f"\n{category}:")
            for skill in skills:
                print(f"- {skill}")
        
        # Generate a tailored resume
        logger.info("\nGenerating tailored resume...")
        output_path = OUTPUT_DIR / f"tailored_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        tailor.create_tailored_resume(tailored_data, str(output_path))
        
        logger.info(f"\n✅ Tailored resume generated successfully: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"❌ Error during resume tailoring test: {str(e)}", exc_info=True)
        return False

if __name__ == "__main__":
    test_resume_tailoring()
