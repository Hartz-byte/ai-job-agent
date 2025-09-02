import logging
from datetime import datetime
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import after logging is configured
from src.parsers.resume_parser import parse_resume  # noqa: E402
from src.generators.services.resume_tailor import ResumeTailor  # noqa: E402

# Create output directory if it doesn't exist
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

def create_test_job():
    """Create a test job with sample requirements."""
    class TestJob:
        def __init__(self):
            self.title = "AI/ML Engineer"
            self.company = "Tech Innovations Inc."
            self.location = "Remote (India)"
            self.url = "https://example.com/job/ai-ml-engineer-123"
            self.description = """
            We are looking for a skilled AI/ML Engineer with experience in deep learning and MLOps.
            
            Key Responsibilities:
            - Design and implement machine learning models for various NLP and computer vision tasks
            - Develop and optimize end-to-end ML pipelines for production deployment
            - Work with large language models (LLMs) and fine-tune them for specific use cases
            - Implement MLOps best practices for model training, deployment, and monitoring
            - Collaborate with cross-functional teams to integrate AI solutions into products
            
            Requirements:
            - Strong experience with Python, PyTorch, and TensorFlow
            - Hands-on experience with deep learning for NLP and/or computer vision
            - Familiarity with MLOps tools (Docker, Kubernetes, MLflow, etc.)
            - Experience with cloud platforms (AWS, GCP, or Azure)
            - Strong understanding of machine learning algorithms and statistics
            - Experience with version control (Git) and CI/CD pipelines
            
            Nice to Have:
            - Experience with large language models (LLMs) and prompt engineering
            - Knowledge of frontend frameworks (Streamlit, React, etc.)
            - Experience with big data technologies (Spark, Dask, etc.)
            - Contributions to open-source ML projects
            - Published research papers or technical blog posts
            """
    return TestJob()

def generate_test_resume():
    """Generate a test resume document with sample content."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
    
    doc = Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add header with name and contact info
    header = doc.add_paragraph()
    header_run = header.add_run("HARSH GUPTA")
    header_run.bold = True
    header_run.font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run("Gurugram, Haryana • +91-7081771202 • harshmail281199@gmail.com")
    contact_run.add_break()
    contact_run.add_text("LinkedIn: https://www.linkedin.com/in/harsh-gupta-dev/ • GitHub: https://github.com/Hartz-byte • Kaggle: https://kaggle.com/hartzbyte")
    
    # Add summary section
    doc.add_heading("SUMMARY", level=2)
    doc.add_paragraph("""
    AI/ML Engineer with hands-on experience designing, training, and scaling deep-learning solutions for NLP, 
    computer vision, and generative AI. Skilled in Python, PyTorch, TensorFlow, and modern MLOps (Docker, AWS, FastAPI). 
    Proven at building end-to-end pipelines, optimizing model performance, and integrating large-language-model (LLM) 
    services into production applications. Agile collaborator who bridges research prototypes and real-world impact.
    """.strip())
    
    # Add experience section
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
    
    # Experience 1 - ML Intern
    exp1 = doc.add_paragraph()
    exp1.add_run("ML Intern").bold = True
    exp1.add_run(" | Tensaw Technologies").bold = False
    exp1.add_run(" | ")
    exp1.add_run("June 2025 - August 2025").italic = True
    
    bullets = [
        "Developed and deployed ML models for payment-risk scoring and anomaly detection in healthcare-fintech datasets using Python, scikit-learn, and TensorFlow",
        "Executed comprehensive data preprocessing, feature engineering, and model evaluation workflows",
        "Presented data-driven insights and model performance metrics to stakeholders in weekly sprint reviews",
        "Built and maintained Dockerized ML pipelines on AWS, ensuring HIPAA-compliant data governance standards"
    ]
    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(bullet)
    
    # Experience 2 - Full Stack Engineer
    doc.add_paragraph()  # Add space between experiences
    exp2 = doc.add_paragraph()
    exp2.add_run("Full Stack Engineer").bold = True
    exp2.add_run(" | Sinqlarity").bold = False
    exp2.add_run(" | ")
    exp2.add_run("January 2023 - March 2025").italic = True
    
    bullets = [
        "Developed and shipped scalable MERN-stack (MongoDB, Express.js, React.js, Node.js) web applications",
        "Optimized backend APIs and implemented secure authentication and payment processing flows",
        "Designed and maintained both relational and NoSQL database schemas for optimal performance",
        "Integrated RESTful AI microservices into web applications, enhancing functionality with machine learning capabilities",
        "Implemented CI/CD pipelines using GitHub Actions, reducing deployment time by 25%",
        "Mentored junior interns on software development best practices and code reviews"
    ]
    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(bullet)
    
    # Add skills section with better formatting
    doc.add_heading("TECHNICAL SKILLS", level=2)
    
    skills = {
        "Programming & Data": ["Python", "Pandas", "NumPy", "SQL", "JavaScript", "TypeScript"],
        "AI/ML & Deep Learning": ["PyTorch", "TensorFlow", "Keras", "Stable Diffusion", "CNN", "RNN", 
                                 "LSTM", "Transformers", "LLM fine-tuning", "Hugging Face", "scikit-learn"],
        "MLOps & Cloud": ["FastAPI", "Docker", "AWS (EC2, S3)", "Git", "CI/CD", "Microservices", 
                         "GPU optimization", "REST APIs"],
        "Frontend & Databases": ["Streamlit", "React.js", "Node.js", "Express.js", "MongoDB", 
                               "PostgreSQL", "Redis"]
    }
    
    # Create a two-column table for skills
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add skills in two columns
    col1, col2 = table.rows[0].cells
    col1_para = col1.paragraphs[0]
    col2_para = col2.paragraphs[0]
    
    # Add skills to columns (alternating between columns for better space usage)
    for i, (category, items) in enumerate(skills.items()):
        target = col1_para if i % 2 == 0 else col2_para
        target.add_run(f"{category}:").bold = True
        target.add_run(" " + ", ".join(items) + "\n")

    
    # Add education section
    doc.add_heading("EDUCATION", level=2)
    
    # Education 1
    edu1 = doc.add_paragraph()
    edu1.add_run("B.Tech in Computer Science").bold = True
    edu1.add_run(" | Graphic Era Deemed to be University | ")
    edu1.add_run("2019 - 2023").italic = True
    
    # Add certifications
    doc.add_heading("CERTIFICATIONS", level=2)
    certs = [
        "Machine Learning Specialization - Coursera (Andrew Ng)",
        "Deep Learning Specialization - Coursera (Andrew Ng)",
        "Natural Language Processing Specialization - Coursera"
    ]
    for cert in certs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(cert)
    
    # Add research/publications section
    doc.add_heading("RESEARCH PUBLICATIONS", level=2)
    pubs = [
        "Transformers Beyond NLP: A Survey on Vision, Speech, and Multimodal Applications",
        "A Survey on Hyperparameter Tuning, Regularization, and Optimization in Deep Neural Networks"
    ]
    for pub in pubs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(pub)
    
    # Save the test resume
    test_resume_path = OUTPUT_DIR / "test_resume.docx"
    doc.save(test_resume_path)
    return test_resume_path

def test_resume_tailoring():
    """Test the resume tailoring functionality."""
    try:
        logger.info("Starting resume tailoring test...")
        
        # Generate a test resume if it doesn't exist
        test_resume_path = OUTPUT_DIR / "test_resume.docx"
        if not test_resume_path.exists():
            logger.info("Generating test resume...")
            test_resume_path = generate_test_resume()
        
        # Parse the test resume
        logger.info("Parsing test resume...")
        profile = parse_resume(str(test_resume_path))
        
        # Create a test job
        test_job = create_test_job()
        
        # Initialize the resume tailor
        logger.info("Initializing ResumeTailor...")
        tailor = ResumeTailor(profile, test_job)
        
        # Generate tailored content
        logger.info("Generating tailored content...")
        tailored_data = tailor.generate_tailored_content()
        
        # Print tailored content
        logger.info("\n=== TAILORED SUMMARY ===")
        print(tailored_data.summary)
        
        logger.info("\n=== TAILORED SKILLS ===")
        for category, skills in tailored_data.technical_skills.items():
            print(f"\n{category}:")
            for skill in skills:
                print(f"- {skill}")
        
        # Generate a tailored resume
        logger.info("\nGenerating tailored resume...")
        output_path = OUTPUT_DIR / f"tailored_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        tailor.create_tailored_resume(tailored_data, str(output_path))
        
        logger.info(f"\n✅ Tailored resume generated successfully: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"❌ Error during resume tailoring test: {str(e)}", exc_info=True)
        return False

if __name__ == "__main__":
    test_resume_tailoring()
